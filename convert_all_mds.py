import os
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from fpdf.enums import XPos, YPos

def clean_text(text):
    replacements = {
        "—": "-", "–": "-", "“": '"', "”": '"', "‘": "'", "’": "'",
        "🛠️": "[Fix]", "🛠": "[Fix]", "📘": "[Guide]", "📌": "[Index]",
        "🌐": "[Web]", "📶": "[LAN]", "🚀": "[Go]", "⚡": "[Fast]",
        "🔥": "[Alert]", "✅": "[OK]", "❌": "[Error]", "•": "-", "…": "..."
    }
    cleaned = text
    for original, sub in replacements.items():
        cleaned = cleaned.replace(original, sub)
    return "".join(c for c in cleaned if ord(c) < 256)

class StyledPDF(FPDF):
    def __init__(self, title_text):
        super().__init__()
        self.title_text = title_text
        self.set_left_margin(20)
        self.set_right_margin(20)
        self.set_top_margin(35)
        self.set_auto_page_break(auto=True, margin=25)
        
    def header(self):
        if self.page_no() == 1:
            return
        self.set_text_color(100, 116, 139)
        self.set_font('Helvetica', 'I', 8.5)
        self.set_y(14)
        self.cell(0, 5, self.title_text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.set_draw_color(226, 232, 240)
        self.set_line_width(0.3)
        self.line(20, 20, 190, 20)
        self.set_top_margin(28)
        self.set_y(24)

    def footer(self):
        if self.page_no() == 1:
            return
        self.set_y(-18)
        self.set_draw_color(226, 232, 240)
        self.set_line_width(0.3)
        self.line(20, 280, 190, 280)
        self.set_font('Helvetica', 'I', 8.5)
        self.set_text_color(148, 163, 184)
        self.cell(0, 10, f'Page {self.page_no()} | System Documentation', align='C')

def build_docx_and_pdf(md_filename, output_base, main_title, subtitle):
    print(f"[*] Processing {md_filename}...")
    if not os.path.exists(md_filename):
        print(f"[-] {md_filename} not found!")
        return

    with open(md_filename, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create Word Doc
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style Word Doc
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(51, 65, 85)

    style_title = doc.styles['Title']
    style_title.font.name = 'Segoe UI'
    style_title.font.size = Pt(28)
    style_title.font.bold = True
    style_title.font.color.rgb = RGBColor(15, 23, 42)

    # DOCX Cover Page
    doc.add_paragraph('\n\n\n\n\n\n')
    t_para = doc.add_paragraph(main_title.upper(), style='Title')
    t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_para = doc.add_paragraph(subtitle, style='Subtitle')
    s_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n\n\n\n\n\n\n\n\n\n\n\n')
    doc.add_page_break()

    # Create PDF
    pdf = StyledPDF(f"Rescue Operations System (ARDMS) - {main_title}")
    pdf.add_page()

    # PDF Cover Page
    pdf.set_fill_color(15, 23, 42)
    pdf.rect(0, 0, 210, 55, 'F')
    pdf.set_text_color(255, 255, 255)
    pdf.set_font('Helvetica', 'B', 15)
    pdf.set_y(15)
    pdf.cell(0, 7, 'RESCUE OPERATIONS SYSTEM (ARDMS)', new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
    pdf.set_font('Helvetica', '', 10)
    pdf.cell(0, 5, 'Advanced Rescue & Disaster Management Suite', new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')

    pdf.set_y(80)
    pdf.set_text_color(15, 23, 42)
    pdf.set_font('Helvetica', 'B', 20)
    pdf.cell(0, 10, main_title.upper(), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
    pdf.ln(2)
    pdf.set_text_color(29, 78, 216)
    pdf.set_font('Helvetica', 'B', 12)
    pdf.cell(0, 7, subtitle, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
    pdf.add_page()

    # Parse and write
    in_code_block = False
    for line in lines:
        raw_line = line.strip('\n')
        line_clean = clean_text(raw_line).strip()
        
        if line_clean.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            # Word Code Block
            p_code = doc.add_paragraph()
            p_code.paragraph_format.left_indent = Inches(0.4)
            run = p_code.add_run(line_clean)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            
            # PDF Code Block
            pdf.set_font("Courier", "B", 8)
            pdf.set_text_color(30, 41, 59)
            pdf.cell(0, 4.5, line_clean, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            continue

        if not line_clean:
            pdf.ln(3)
            continue

        # Header 1
        if line_clean.startswith('# '):
            title = line_clean[2:]
            doc.add_heading(title, level=1)
            
            pdf.ln(6)
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(14, 165, 233)
            pdf.multi_cell(0, 10, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)
            
        # Header 2
        elif line_clean.startswith('## '):
            title = line_clean[3:]
            doc.add_heading(title, level=2)
            
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(30, 41, 59)
            pdf.multi_cell(0, 8, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)

        # Header 3
        elif line_clean.startswith('### '):
            title = line_clean[4:]
            doc.add_heading(title, level=3)
            
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(71, 85, 105)
            pdf.multi_cell(0, 6, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)

        # Bullet List
        elif line_clean.startswith('* ') or line_clean.startswith('- '):
            bullet_text = line_clean[2:]
            
            # Word Bullet
            p_bullet = doc.add_paragraph(style='List Bullet')
            parts = bullet_text.split('**')
            for idx, part in enumerate(parts):
                run = p_bullet.add_run(part)
                if idx % 2 == 1:
                    run.bold = True
                    
            # PDF Bullet
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(51, 65, 85)
            pdf.write(5.5, "- ")
            parts = bullet_text.split('**')
            for idx, part in enumerate(parts):
                if idx % 2 == 1:
                    pdf.set_font('Helvetica', 'B', 10)
                else:
                    pdf.set_font('Helvetica', '', 10)
                pdf.write(5.5, part)
            pdf.ln(6)

        # Numbered List
        elif re.match(r'^\d+\.\s', line_clean):
            match = re.match(r'^(\d+\.\s)(.*)', line_clean)
            num_prefix = match.group(1)
            num_text = match.group(2)
            
            # Word List
            p_num = doc.add_paragraph()
            p_num.paragraph_format.left_indent = Inches(0.2)
            p_num.add_run(num_prefix).bold = True
            parts = num_text.split('**')
            for idx, part in enumerate(parts):
                run = p_num.add_run(part)
                if idx % 2 == 1:
                    run.bold = True
                    
            # PDF List
            pdf.set_font('Helvetica', 'B', 10)
            pdf.set_text_color(51, 65, 85)
            pdf.write(5.5, num_prefix)
            pdf.set_font('Helvetica', '', 10)
            parts = num_text.split('**')
            for idx, part in enumerate(parts):
                if idx % 2 == 1:
                    pdf.set_font('Helvetica', 'B', 10)
                else:
                    pdf.set_font('Helvetica', '', 10)
                pdf.write(5.5, part)
            pdf.ln(6)

        # Standard Paragraph
        else:
            # Word Paragraph
            p_normal = doc.add_paragraph()
            p_normal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parts = line_clean.split('**')
            for idx, part in enumerate(parts):
                run = p_normal.add_run(part)
                if idx % 2 == 1:
                    run.bold = True
                    
            # PDF Paragraph
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(51, 65, 85)
            parts = line_clean.split('**')
            for idx, part in enumerate(parts):
                if idx % 2 == 1:
                    pdf.set_font('Helvetica', 'B', 10)
                else:
                    pdf.set_font('Helvetica', '', 10)
                pdf.write(5.5, part)
            pdf.ln(6.5)

    doc.save(f"{output_base}.docx")
    pdf.output(f"{output_base}.pdf")
    print(f"[+] Successfully created {output_base}.docx and {output_base}.pdf")

if __name__ == '__main__':
    build_docx_and_pdf("DEPLOYMENT_GUIDE.md", "DEPLOYMENT_GUIDE", "Localhost Deployment Guide", "Deployment and System Setup Specification Handbooks")
    build_docx_and_pdf("TECH_STACK.md", "TECH_STACK", "System Technical Stack", "Technical Architecture and Implementation Guide")
    build_docx_and_pdf("PROJECT_WORKFLOW.md", "PROJECT_WORKFLOW", "Project Workflow & Operational Master Guide", "Project Flow and Operational Procedures")
