import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # --- MARGINS SETUP (STRICT 1.0 INCH) ---
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # --- STYLE SETUP ---
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Segoe UI'
    font_normal.size = Pt(10.5)
    font_normal.color.rgb = RGBColor(51, 65, 85) # Slate 700

    style_title = doc.styles['Title']
    font_title = style_title.font
    font_title.name = 'Segoe UI'
    font_title.size = Pt(28)
    font_title.bold = True
    font_title.color.rgb = RGBColor(15, 23, 42) # Slate 900
    
    style_h1 = doc.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Segoe UI'
    font_h1.size = Pt(16)
    font_h1.bold = True
    font_h1.color.rgb = RGBColor(14, 165, 233) # Sky 500

    style_h2 = doc.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Segoe UI'
    font_h2.size = Pt(13)
    font_h2.bold = True
    font_h2.color.rgb = RGBColor(30, 41, 59) # Slate 800

    style_h3 = doc.styles['Heading 3']
    font_h3 = style_h3.font
    font_h3.name = 'Segoe UI'
    font_h3.size = Pt(11)
    font_h3.bold = True
    font_h3.color.rgb = RGBColor(71, 85, 105) # Slate 600
    
    # --- TITLE PAGE ---
    doc.add_paragraph('\n\n\n\n\n\n')
    title = doc.add_paragraph('ARDMS SYSTEM MANUAL', style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Advanced Rescue & Disaster Management System\nTechnical Operations Manual & System Workflows', style='Subtitle')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(100, 116, 139) # Slate 500
    doc.add_paragraph('\n\n\n\n\n\n\n\n\n\n\n\n\n')
    doc.add_page_break()

    def h1(text): 
        doc.add_heading(text, level=1)
    def h2(text): 
        doc.add_heading(text, level=2)
    def h3(text): 
        doc.add_heading(text, level=3)
    
    def p(text): 
        p_item = doc.add_paragraph(text)
        p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p_item

    def b(text): 
        p_item = doc.add_paragraph(text, style='List Bullet')
        p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p_item

    def flowchart(lines):
        p_flow = doc.add_paragraph()
        p_flow.paragraph_format.left_indent = Inches(0.2)
        p_flow.paragraph_format.right_indent = Inches(0.2)
        p_flow.paragraph_format.space_before = Pt(6)
        p_flow.paragraph_format.space_after = Pt(6)
        
        full_text = ""
        for line in lines:
            full_text += line + "\n"
        
        run = p_flow.add_run(full_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        run.bold = True

    # --- CONTENT ---
    h1("A. Project Overview & Abstract")
    p("The Advanced Rescue & Disaster Management System (ARDMS) is a complete digital solution designed to save lives during critical emergencies. It connects people in danger directly with nearby rescue teams and a central command center. By utilizing real-time GPS tracking, instant SOS alerts, and smart task grouping, ARDMS eliminates the confusion and delays typically seen during natural disasters or medical emergencies. This project provides a Web Admin Dashboard for operators, a Public App for victims to ask for help, and a Rescuer App for first responders to navigate exactly where they are needed.")

    h1("B. Introduction")
    p("When a disaster strikes - like a flood, earthquake, or severe medical emergency - every single second matters. Unfortunately, traditional phone calls to emergency numbers can get jammed, and explaining your exact location can be difficult when you are in a panic.")
    p("ARDMS solves this problem by using the technology already in our pockets: smartphones. With ARDMS, a person just has to press a single SOS button. Instantly, their exact GPS coordinates, the type of emergency, and their details are sent to a live Web Dashboard. An operator immediately sees this on a map and dispatches a registered Rescuer to that exact spot. The system acts as a digital bridge between those who need help and those who can provide it.")

    h1("C. System Architecture & Tech Stack")
    p("The system runs on a modern, real-time technology stack. It uses React Native for the mobile apps, React.js for the Web Dashboard, and a Node.js/SQLite backend for the database. Everything communicates instantly using WebSockets.")
    b("Database Core: SQLite3 running in WAL (Write-Ahead Logging) mode to support concurrent transactions.")
    b("Backend Server: Express.js REST API with integrated WebSockets for real-time telemetry streaming.")
    b("Web Admin Center: HTML5 / Vanilla JS Single Page Dashboard styled with premium custom dark mode and glassmorphism, rendering dynamic incident markers via Leaflet.js.")
    b("Mobile Applications: HTML5 Progressive Web App templates (Field Rescuer and Public SOS Apps) leveraging native GPS modules and offline AsyncStorage cache queues.")

    doc.add_page_break()
    h1("D. Master Workflow Overview")
    p("Below is the complete overview diagram tracing the flow of an emergency. It traces the signal from the Public App triggering an SOS, to the Admin Dashboard dispatching it, directly to the Rescuer App accepting it.")
    try:
        doc.add_picture('master_flow.png', width=Inches(6.5))
    except Exception: pass

    p("1. Victim opens Public App and taps 'Medical Emergency'.")
    p("2. App waits 3 seconds, captures GPS, and POSTs data to Server.")
    p("3. Server saves to Database and alerts Web Admin.")
    p("4. Admin sees a flashing Red Alert on their map.")
    p("5. Admin clicks the alert, selects Rescuer 'John', and clicks 'Assign'.")
    p("6. John's Rescuer App rings. John clicks 'Accept'.")
    p("7. John follows the blue navigation line on his map to the victim.")
    p("8. Victim sees John's vehicle icon moving toward them on their screen.")
    p("9. John arrives, helps the victim, clicks 'Complete', and the alert turns green on the Admin dashboard.")

    doc.add_page_break()
    
    # --- MASTER LIST OF OPERATIONS ---
    h1("E. Master List of Operations")
    p("The ARDMS system operates through a set of interconnected workflows. Below is a structured list of every operation involved in the ecosystem, divided by platform and functional layer:")

    h2("1. Triggering Requests (Public Citizen App Operations)")
    b("Operation 1.1: Multi-Category Emergency Distress Broadcasting - Allows citizens to select specific emergency types (Medical, Flood, Fire, General Support) which triggers high-accuracy GPS capture, initiates a 3-second buffer cooldown timer, and uploads the payload.")
    b("Operation 1.2: Real-Time Incident Telemetry Sync & Map Plotting - Listens to WebSocket broadcasts from the server containing the responder's moving coordinates to animate a vehicle icon.")
    b("Operation 1.3: User-Initiated Distress SOS Resolution & Abort Flow - Safely cancels an active request, sweeps database tables, and returns the citizen to the main home UI.")
    b("Operation 1.4: Offline Network Resilience & AsyncStorage Synchronization - Caches SOS payloads in local memory during poor coverage, bulk uploading them immediately on recovery.")

    h2("2. Rescuer Operations (Rescuer Mobile Field App Operations)")
    b("Operation 3.1: Instant Mission Incoming Handshake Alert & Wake Lock - Alerts responders via loud alarms, vibrations, device wake locks, and yellow flashing card overlays.")
    b("Operation 3.2: Double-Lock Mission Acceptance & Real-Time GPS Activation - Updates mission command state to 'in-progress' in the database and triggers GPS tracking loops.")
    b("Operation 3.3: Geographic Convex Hull Border Plotting & Zone Boundary Visualization - Renders a translucent purple outline around grouped tasks on the map with coordinate safety fallbacks.")
    b("Operation 3.4: WebSocket Background Geolocation Telemetry Streaming - Automatically streams high-accuracy coordinates to `/api/telemetry` over WebSockets every few seconds.")
    b("Operation 3.5: Physical Incident Resolution & Database State Finalization - Completes the task, resolving the active dashboard line and cascading states to all child items.")
    b("Operation 3.6: Persistent Offline History Logging & Session Preservation - Archives completed/declined tasks in local memory, preserving credentials during selective cache washes.")

    h2("3. Web System Operations (Web Admin Dashboard Operations)")
    b("Operation 2.1: Multi-Channel Real-Time Incident Triage & Leaflet Mapping - Renders real-time incidents on Leaflet maps, featuring split columns for image preview lightboxes and inline WAV players.")
    b("Operation 2.2: Proximity-Based Individual Dispatch & Mission Initialization - Tracks active rescuers, computes proximity lines, ranks rescuers by distance, and dispatches the task.")
    b("Operation 2.3: Multi-Incident Tactical Grouping & Convex Hull Boundary Polygon Plotting - Bypasses auto-refresh race conditions using persistent `window.bulkSelectedIds` states to dispatch supply clusters.")
    b("Operation 2.4: Active Route Polyline Memory Management & Map Cleanup - Auto-erases route polylines upon completion or rejection to avoid visual map clutter.")
    b("Operation 2.5: Dynamic In-Page Crew Re-allocation & Real-Time Redirection - Re-routes active responders mid-flight using in-page reassign modal pop-ups without page reloads.")

    h2("4. Core Backend & Database Operations")
    b("Operation 4.1: Database Write-Ahead Logging (WAL) Mode Initialization - Enables SQLite WAL mode on boot to allow concurrent, non-locking reads and writes during incident spikes.")
    b("Operation 4.2: Dynamic Schema Migration & Verification - Validates database table parameters on startup, dynamically injecting new fields like `audio_url` without data loss.")
    b("Operation 4.3: Base64 Audio SOS Decode and Directory Ingestion - Decodes base64 wave streams, saves binary audio files in folders, and links their path to the SQLite record.")
    b("Operation 4.4: WebSocket Dynamic Multiroom Message Routing - Handles websocket connections, channels coordinates, and maps updates to correct room pipelines.")

    h2("5. Network Deployment Operations")
    b("Operation 5.1: Private Network-in-a-Box (NIB) Cellular Base Station Setup - Deploys private cellular SDR transceivers and programs private USIM cards for local coverage.")
    b("Operation 5.2: Tailscale Private Mesh VPN Tunneling - Bypasses Carrier-Grade NAT (CGNAT) using virtual private networks to link remote devices over standard cell networks.")
    b("Operation 5.3: Windows Defender Inbound Firewall Access Rule Setup - Creates PowerShell inbound firewall rules on port 3001 to prevent dynamic LAN connection drops.")

    doc.add_page_break()

    # --- DEDICATED OPERATIONS AND WORKFLOWS ---
    h1("F. Detailed System Operations & Working Flows")
    p("This section describes each operation in detail, detailing the technical description, step-by-step workflow, high-tech flow diagrams, and user interface layouts.")

    # SECTION F.1: PUBLIC APP
    h2("F.1 Public Mobile Support App Operations")
    p("The Public Mobile App allows citizens to broadcast emergency signals, track dispatched rescue units, and queue submissions during poor network conditions.")

    # OP 1.1
    h3("Operation 1.1: Multi-Category Emergency Distress Broadcasting")
    p("Description: Allows citizens to select specific emergency types (Medical, Flood, Fire, General Support) which triggers a high-accuracy GPS capture and uploads a distress event. Initiates a 3-second countdown delay buffer to allow user cancellation in case of accidental clicks.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [Citizen Home UI] -> Select Emergency Type (Fire/Medical) |",
        "|         |                                                  |",
        "|  [3s Cooldown Timer] --(Cancel Tapped?)-> [Abort SOS]      |",
        "|         |                                                  |",
        "|  [Lock High-Accuracy GPS Coordinates]                      |",
        "|         |                                                  |",
        "|  [POST Distress Payload to /api/requests]                  |",
        "|         |                                                  |",
        "|  [WebSocket Broadcast to Web Admin] -> [Live Map UI]       |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Open the ARDMS Public App on the mobile device.")
    b("Step 2: Identify the nature of the emergency on the home screen.")
    b("Step 3: Tap the corresponding emergency button (e.g., 'Medical Emergency' or 'Flood Rescue').")
    b("Step 4: A 3-second countdown banner (#selectionBufferCooldownBanner) appears. Wait or tap 'Cancel SOS' if triggered accidentally.")
    
    p("Visual UI Details (Home Screen Layout):")
    p("The mockup image below represents the emergency category selection home screen. It features a modern, accessible interface with warm slate styling. The primary buttons ('Medical Emergency', 'Flood Rescue', 'Fire Emergency') are large and high-contrast, ensuring ease of use for citizens under extreme panic or stress conditions.")
    try:
        doc.add_picture('public_home.png', width=Inches(2.5))
    except Exception: pass
    p("Note: The category selection maps to urgency parameters on the backend (e.g. 'Critical' for Medical/Fire/Flood and 'Normal' for general support).")

    # OP 1.2
    h3("Operation 1.2: Real-Time Incident Telemetry Sync & Map Plotting")
    p("Description: Connects the Public App to the server's WebSocket tracking rooms. The app continuously listens for location updates from the assigned responder and animates a custom map vehicle marker advancing toward the citizen.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [Establish Server WebSocket] -> Join Room 'rescue-room'  |",
        "|         |                                                  |",
        "|  [Listen for Rescuer Telemetry Broadcast]                  |",
        "|         |                                                  |",
        "|  [Extract Live Latitude & Longitude]                       |",
        "|         |                                                  |",
        "|  [Animate Rescuer Marker] -> [Update live ETA & Distance]  |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Once an SOS is active and assigned, the app automatically navigates to the Tracking page.")
    b("Step 2: Verify the map loads showing your location and a blue vehicle marker.")
    b("Step 3: Track the movement of the responder in real time along with the live distance readout.")
    
    p("Visual UI Details (Active Tracking Map UI):")
    p("The mockup image below shows the active tracking map. A red marker plots the citizen's static position, and a moving blue marker displays the rescuer's live telemetry coordinate in real-time, accompanied by estimated arrival metrics and an absolute 'Cancel SOS' button at the bottom.")
    try:
        doc.add_picture('public_sos.png', width=Inches(2.5))
    except Exception: pass
    p("Note: If the network drops during this operation, the app displays a prominent status indicator showing 'Offline - Reconnecting' while retaining the last cached coordinates of the responder.")

    # OP 1.3
    h3("Operation 1.3: User-Initiated Distress SOS Resolution & Abort Flow")
    p("Description: Allows citizens to cancel their SOS alert if the situation resolves or the call was accidental, sending a delete command to the server and cleaning active command states.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [User clicks 'Cancel SOS' Button]                         |",
        "|         |                                                  |",
        "|  [REST DELETE /api/requests/:id/cancel]                    |",
        "|         |                                                  |",
        "|  [Database Sweeps State] -> [WS Broadcast to Web Admin]    |",
        "|         |                                                  |",
        "|  [Clear active session tokens] -> [Redirect to Home UI]    |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Tap the 'Cancel SOS' button at the bottom of the tracking screen.")
    b("Step 2: Confirm the cancellation prompt on the pop-up modal.")
    b("Step 3: The app immediately returns to the home page, clearing active alert flags.")

    # OP 1.4
    h3("Operation 1.4: Offline Network Resilience & AsyncStorage Synchronization")
    p("Description: Implements local offline caches. When the system detects a network drop, it saves the SOS report in local device memory and dynamically schedules queue pushes upon internet recovery, preventing data loss.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [Trigger SOS Alert] -> [Detect Network Outage]            |",
        "|         |                                                  |",
        "|  [Write distress payload to AsyncStorage 'offline_queue']  |",
        "|         |                                                  |",
        "|  [Render Offline Alert Banner] -> [Monitor Background ping]|",
        "|         |                                                  |",
        "|  [Network Restored] -> [REST Bulk POST queued SOS events]  |",
        "|         |                                                  |",
        "|  [Clear Local Queue] -> [Restore active WebSockets]        |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Trigger an emergency alert while offline.")
    b("Step 2: A yellow notification bar will alert you that your request is safely stored in local cache.")
    b("Step 3: When cellular network is re-established, the app automatically synchronizes and updates the database.")

    doc.add_page_break()

    # SECTION F.2: WEB ADMIN
    h2("F.2 Web Admin Dashboard Operations")
    p("The Command Center Web Dashboard gives dispatcher teams comprehensive mapping controls, tactical group building tools, and real-time crew assignment utilities.")

    # OP 2.1
    h3("Operation 2.1: Multi-Channel Real-Time Incident Triage & Leaflet Mapping")
    p("Description: Integrates raw WebSocket packets with Leaflet.js map layers, showing flashing red incident markers for urgent requests, while rendering separate media columns for IMAGE attachments and inline play/download AUDIO WAV players.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [WS Connection Ready] -> [Capture Incoming SOS Alert]     |",
        "|         |                                                  |",
        "|  [Query DB Details] -> [Format & Parse Media Uploads]      |",
        "|         |                                                  |",
        "|  [Plot Flashing Red Leaflet Marker] -> [Render split media]|",
        "|         |                                                  |",
        "|  [Play Sound Indicator] -> [Populate Triage Grid Columns]  |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Open the ARDMS Command Center portal in a browser.")
    b("Step 2: Review the map grid displaying active rescue markers.")
    b("Step 3: Access incoming media reports directly via the split columns (thumbnails & inline WAV player).")
    
    p("Visual UI Details (Command Dashboard & Proximity Dispatch):")
    p("The mockup image below represents the Web Admin Command Center dispatch panel. The central map displays active incidents (red markers) and responder vectors (blue markers). Clicking a marker opens the sidebar dispatch UI, allowing operators to match the task with a proximity-ranked rescuer.")
    try:
        doc.add_picture('admin_assign.png', width=Inches(5.0))
    except Exception: pass
    p("Note: The dispatch dropdown lists rescuers in real-time, sorted by proximity to ensure fast arrivals.")

    # OP 2.2
    h3("Operation 2.2: Proximity-Based Individual Dispatch & Mission Initialization")
    p("Description: Maps individual rescue tasks. When an incident is selected, the system queries the coordinates of all active online rescuers, calculates direct GPS line distance, ranks them, and dispatches the mission to the closest available officer.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [Operator clicks Pin] -> [Query all active rescuers]     |",
        "|         |                                                  |",
        "|  [Calculate Proximity Metric (GPS coords distance)]        |",
        "|         |                                                  |",
        "|  [Sort Rescuer dropdown dynamically by shortest distance]  |",
        "|         |                                                  |",
        "|  [Click Dispatch] -> [POST /api/commands] -> [WS to Mobile]|",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Select an active emergency marker on the Leaflet map.")
    b("Step 2: In the dispatch panel, review the proximity-sorted list of nearby rescuers.")
    b("Step 3: Select the optimal unit and click 'Dispatch Unit' to send the mission details.")

    # OP 2.3
    h3("Operation 2.3: Multi-Incident Tactical Grouping & Convex Hull Boundary Polygon Plotting")
    p("Description: Bundles multiple incident markers in a disaster zone into a single group. Stores checkboxes inside `window.bulkSelectedIds` globally so background auto-refresh runs do not clear the selection. Computes a convex hull bounding box around selected markers and dispatches a single heavy-duty squad with a unified multi-waypoint zone route.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [Operator checks multiple emergency checkboxes]           |",
        "|         |                                                  |",
        "|  [Cache selections inside window.bulkSelectedIds]          |",
        "|         |                                                  |",
        "|  [Open Bulk Group View] -> [Compute Convex Hull Polygon]   |",
        "|         |                                                  |",
        "|  [Specify supply list] -> [Select heavy squad team]        |",
        "|         |                                                  |",
        "|  [POST /api/commands/group] -> [Draw boundary hull on map] |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Check the selection boxes next to multiple emergencies in the triage board.")
    b("Step 2: Click 'Group & Assign' to create a Tactical Cluster.")
    b("Step 3: Review the computed convex hull polygon bounding box on the selection map.")
    b("Step 4: Specify necessary supplies (e.g. water boxes, emergency blankets).")
    b("Step 5: Select the target crew and click 'CONFIRM & INITIALIZE GROUP MISSION'.")
    
    p("Visual UI Details (Tactical Convex Hull Grouping):")
    p("The mockup image below displays the tactical group creation view. A bounding box polygon is computed and plotted around a cluster of multiple flood victims in Sector B. The sidebar transitions to the 'Grouped Supplies Request' panel for dispatching a single heavy-duty responder.")
    try:
        doc.add_picture('admin_group.png', width=Inches(5.0))
    except Exception: pass
    p("Note: This bulk method avoids redundant dispatches and reduces radio/data traffic for multiple incidents in close proximity.")

    # OP 2.4
    h3("Operation 2.4: Active Route Polyline Memory Management & Map Cleanup")
    p("Description: Manages map memory. When a task is completed, resolved, or declined, the system calls `cleanActiveRouteLayers()` to identify and remove matching polyline route graphics from the map canvas, keeping the view clean.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [WS Alert: Command status resolved/completed/declined]   |",
        "|         |                                                  |",
        "|  [Trigger cleanActiveRouteLayers(commandId)]               |",
        "|         |                                                  |",
        "|  [Locate matching Leaflet L.polyline layers in cache]      |",
        "|         |                                                  |",
        "|  [Call polyline.remove() to erase lines from map canvas]   |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: The map shows blue dashed lines indicating active responder paths.")
    b("Step 2: When an officer marks a task complete or cancels, the line immediately vanishes from the screen.")

    # OP 2.5
    h3("Operation 2.5: Dynamic In-Page Crew Re-allocation & Real-Time Redirection")
    p("Description: Allows operators to redirect a squad mid-flight. Clicking 'Re-assign' accesses the command ID and triggers the re-assignment modal, sending a redirection handshake to both the original and new mobile apps.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [Click 'Re-assign' next to active command cluster]        |",
        "|         |                                                  |",
        "|  [Locate active Command ID and load openReassignModal()]   |",
        "|         |                                                  |",
        "|  [Select new crew in dropdown] -> [POST /api/reassign]     |",
        "|         |                                                  |",
        "|  [Original app gets abort] -> [New app receives incoming]  |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Click the 'RE-ASSIGN' button next to an active command.")
    b("Step 2: Choose the new responder and click 'Update Assignment'.")
    b("Step 3: Confirm the redirection update; the original task is reassigned without reloading the page.")

    doc.add_page_break()

    # SECTION F.3: RESCUER APP
    h2("F.3 Rescuer Mobile Field App Operations")
    p("The Rescuer Mobile App provides emergency responders with continuous routing instructions, incident details, and direct synchronization with the Command Center.")

    # OP 3.1
    h3("Operation 3.1: Instant Mission Incoming Handshake Alert & Wake Lock")
    p("Description: Triggers high-priority visual and audio alerts when a new mission is dispatched, sounding a loud ringer, pulsing the motor, and waking the device screen so rescuers are notified instantly.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [WS Command broadcast parsed by Rescuer App client]       |",
        "|         |                                                  |",
        "|  [Activate Android Device Wake Lock] -> [Loop loud alarm]  |",
        "|         |                                                  |",
        "|  [Flash high-visibility yellow 'INCOMING MISSION' alert]   |",
        "|         |                                                  |",
        "|  [Render mission distance, ETA, and priority level]        |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Keep the Rescuer App running.")
    b("Step 2: When a task is dispatched, the phone will ring loudly and display the incoming mission screen.")
    b("Step 3: Review the distance, estimated time, and incident type.")
    
    p("Visual UI Details (Incoming Mission Interface):")
    p("The mockup image below represents the high-priority alert flashed on the rescuer's smartphone. The entire card flashes yellow, looping a loud warning sound. The rescuer can tap the large green 'ACCEPT TASK' button or decline, routing the task back.")
    try:
        doc.add_picture('rescuer_accept.png', width=Inches(2.5))
    except Exception: pass
    p("Note: The wake lock prevents the screen from turning off, ensuring notifications are visible even when the device is locked.")

    # OP 3.2
    h3("Operation 3.2: Double-Lock Mission Acceptance & Real-Time GPS Activation")
    p("Description: Locks the rescuer to the task by posting an accept state to the database, switching the app to navigation mode and activating active GPS updates.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [Rescuer taps green 'ACCEPT TASK' button]                |",
        "|         |                                                  |",
        "|  [POST /api/commands/:id/accept] -> Update status in DB    |",
        "|         |                                                  |",
        "|  [Retrieve target coordinates] -> [Initialize GPS tracking]|",
        "|         |                                                  |",
        "|  [Lock active accept states] -> [Navigate to Tracking UI]  |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Tap the green 'ACCEPT TASK' button on the alert card.")
    b("Step 2: The app will update its status and transition to the map view.")

    # OP 3.3
    h3("Operation 3.3: Geographic Convex Hull Border Plotting & Zone Boundary Visualization")
    p("Description: Visualizes group mission zones. The app reads the `missions` list of coordinates and draws a translucent boundary. To prevent crashes or zooming to `[0, 0]` when center coordinates are missing, `startMission()` falls back to using the coordinates of the first mission.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [Receive group command payload]                           |",
        "|         |                                                  |",
        "|  - Center Coordinates Missing? -> Fallback to 1st mission  |",
        "|         |                                                  |",
        "|  [Calculate Convex Hull envelope surrounding all tasks]    |",
        "|         |                                                  |",
        "|  [Draw purple border overlay] -> [Plot markers for childs] |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Accept a grouped tactical mission.")
    b("Step 2: Observe the map zoom to the region, plotting a translucent purple outline around all selected victims.")
    b("Step 3: Tap any victim pin within the boundary to view their specific needs.")
    
    p("Visual UI Details (Active Incident Navigation View):")
    p("The mockup image below represents the active tracking map on the rescuer app. A blue path outlines the direct route to the target location. Proximity distance and turn instructions are visible in the overlay, and a 'MARK COMPLETED' button is placed at the bottom.")
    try:
        doc.add_picture('rescuer_nav.png', width=Inches(2.5))
    except Exception: pass
    p("Note: Geolocation parameters update coordinates in the background via REST, ensuring data consistency.")

    # OP 3.4
    h3("Operation 3.4: WebSocket Background Geolocation Telemetry Streaming")
    p("Description: Streams location updates. When a mission is active, a background thread monitors GPS changes and streams them to the server over WebSockets every few seconds, updating both the Admin and Citizen maps.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [Mission In-Progress] -> [Start background location thread|",
        "|         |                                                  |",
        "|  [Capture continuous device GPS coords (lat, lng)]         |",
        "|         |                                                  |",
        "|  [Send telemetry packet to /api/telemetry over WS]         |",
        "|         |                                                  |",
        "|  [Command center and Citizen client update live vectors]   |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Proceed along the route towards the incident location.")
    b("Step 2: The app automatically streams your current position to the server in the background.")

    # OP 3.5
    h3("Operation 3.5: Physical Incident Resolution & Database State Finalization")
    p("Description: Completes the task. Tapping 'MARK COMPLETED' sends a complete command to the server, which sets the status to completed in the database, cascades the status to any child requests in a group, and archives the event.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [Rescuer arrives at scene & resolves the rescue incident] |",
        "|         |                                                  |",
        "|  [Taps 'MARK COMPLETED' button]                            |",
        "|         |                                                  |",
        "|  [POST /api/commands/:id/complete]                         |",
        "|         |                                                  |",
        "|  [Server cascades state to child requests]                 |",
        "|         |                                                  |",
        "|  [Clear routing layers] -> [Revert back to main feed UI]   |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Upon physical arrival and successful resolution of the rescue task, tap 'MARK COMPLETED'.")
    b("Step 2: Verify that your map is cleared of the route, returning you to the active task feed.")

    # OP 3.6
    h3("Operation 3.6: Persistent Offline History Logging & Session Preservation")
    p("Description: Caches history logs. The app queries specialized history endpoints to show a list of active, completed, and declined tasks. The clear cache function selectively removes history logs without affecting active session tokens.")
    flowchart([
        "+------------------------------------------------------------+",
        "|  [Navigate to History Tab] -> [Fetch complete user history]|",
        "|         |                                                  |",
        "|  [Render all-status logs (completed, declined, active)]    |",
        "|         |                                                  |",
        "|  [Mirror log archive to AsyncStorage for offline backup]   |",
        "|         |                                                  |",
        "|  [Selective cache clear sweeps history, preserves session] |",
        "+------------------------------------------------------------+"
    ])
    p("Step-by-Step Instructions:")
    b("Step 1: Open the Rescuer App and select the 'History' tab on the navigation bar.")
    b("Step 2: Scroll through the history feed to review your completed assignments.")
    b("Step 3: Tap 'Clear History Cache' in settings to clean data logs while remaining logged in.")
    
    p("Visual UI Details (Rescuer History Panel):")
    p("The mockup image below represents the mission history logs screen. Chronologically ordered cards display previous missions. Tapping a card displays complete dispatch logs, response times, and final resolutions.")
    try:
        doc.add_picture('rescuer_history.png', width=Inches(2.5))
    except Exception: pass
    p("Note: The offline cache ensures rescuers can access these logs even in dead network zones.")

    # --- TASK LIFECYCLE ---
    doc.add_page_break()
    h1("G. Task Lifecycle Workflow")
    p("Every emergency requested by the Public App becomes a 'Task' in the ARDMS system. A Task follows a strict operational lifecycle to ensure no emergency is left unresolved.")
    b("State 1: PENDING - The moment an SOS is triggered, a database record is created. It appears as a flashing red alert on the Admin Dashboard.")
    b("State 2: DISPATCHED - Once the Admin selects a Rescuer and clicks 'Dispatch', the payload is pushed to the Rescuer's device.")
    b("State 3: IN-PROGRESS - When the Rescuer hits 'ACCEPT TASK', the status changes. The routing line is drawn, and the ETA begins counting down.")
    b("State 4: COMPLETED - Upon arriving at the scene and resolving the crisis, the Rescuer presses 'MARK COMPLETED'. The task is cleared from the active map and securely logged into the History Database.")

    h1("H. Dynamic Map & Routing Workflow")
    p("The mapping system is the visual core of ARDMS, providing real-time geographical awareness.")
    b("1. Initial Geolocation: The Public App uses native GPS modules to lock onto high-accuracy coordinates (Latitude/Longitude) before transmitting the SOS payload.")
    b("2. Admin Plotting: The Node.js server receives these coordinates and instantly plots a Red Emergency Marker on the Web Dashboard's live map.")
    b("3. Routing Calculation: Upon task assignment, the Rescuer App queries routing APIs to draw a blue tactical path linking their current location to the victim's location.")
    b("4. Live Telemetry Sync: As the Rescuer physically moves, their phone pulses updated coordinates every few seconds via WebSockets. The server reflects this movement, visibly animating the Rescuer's icon advancing along the route on both the Admin's screen and the Victim's screen.")

    h1("I. Autonomous Task Management (AI & Buffer System)")
    p("To eliminate manual delays during high-stress disaster events, the system introduces fully autonomous triage and task distribution.")
    b("1. AI Auto-Assignment: An automated backend process dynamically assesses proximity, workload, and unit types to assign unhandled SOS requests to the best available field rescuer instantly.")
    b("2. Dynamic Reassignment Buffer: The system incorporates a strict fail-safe timeout logic. If an assigned rescuer loses internet connectivity, ignores the task beyond the configured interval, or manually declines the dispatch, the assignment is instantly revoked and auto-routed to the next optimal candidate.")
    b("3. Intelligent Siren State Management: Critical dispatch alerts physically ring the rescuer's device. To prevent debilitating audio loops in the field, the system enforces a strict state-machine that ceases the siren precisely upon task acceptance, decline, or auto-reassignment.")
    b("4. Cache Evasion Techniques: Utilizing strict cache-busting telemetry protocols, the mobile React Native WebViews are forced to bypass aggressive localized caching, guaranteeing absolute synchronization of task states across all units.")

    file_path = os.path.join(os.getcwd(), 'ARDMS_Project_Report_User_Manual_v6.docx')
    doc.save(file_path)
    print(f"Docx saved to {file_path}")

if __name__ == "__main__":
    create_report()
